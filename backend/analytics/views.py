import pandas as pd
import plotly.graph_objects as go
import matplotlib.pyplot as plt
import seaborn as sns
from fpdf import FPDF
import os
import plotly.express as px
import matplotlib
from docx import Document
from docx.shared import Inches
matplotlib.use('Agg')

# =======================
# FUNCIONES DEMOGRÁFICAS
# =======================

# Gráfico de pastel genérico 
def grafico_pie(df, columna, img_path, titulo, leyenda, borde=False):
    df_grouped = df[columna].value_counts()
    colors = plt.cm.Paired.colors[:len(df_grouped)]
    plt.figure(figsize=(4, 4))  
    wedgeprops = {"edgecolor": "black", "linewidth": 0.5} if borde else None
    wedges, texts, autotexts = plt.pie(
        df_grouped,
        labels=df_grouped.index,
        autopct='%1.1f%%',
        startangle=90,
        colors=colors,
        wedgeprops=wedgeprops
    )
    plt.title(titulo, fontsize=10, pad=10) 
    plt.legend(wedges, df_grouped.index, title=leyenda, loc="center left", bbox_to_anchor=(1, 0.5))
    plt.tight_layout()  
    plt.savefig(img_path, bbox_inches='tight')
    plt.close()

def grafico_tipo_sexo(df, img_path):
    grafico_pie(df, "2. SEXO", img_path, "TIPO DE SEXO", "SEXO", borde=True)

def grafico_tipo_vivienda(df, img_path):
    grafico_pie(df, "9. TIPO DE VIVIENDA", img_path, "TIPO DE VIVIENDA", "TIPO DE VIVIENDA")

def grafico_personas_a_cargo(df, img_path):
    grafico_pie(df, "10. NUMERO DE PERSONAS QUE DEPENDEN ECONOMICAMENTE DE USTED", img_path,
                "Personas Económicamente a Cargo", "# Personas a Cargo")

def grafico_tipo_contrato(df, img_path):
    grafico_pie(df, "17. TIPO DE CONTRATO ACTUAL", img_path, "TIPO DE CONTRATO", "TIPO DE CONTRATO")

def grafico_estrato(df, img_path):
    grafico_pie(df, "8. ESTRATO SOCIAL", img_path, "TIPO ESTRATO", "     Tipo de Estrato   ")

# Gráfico de pastel en 3D para estado civil 
def grafico_estadocivil(df, img_path):
    df_grouped = df["4. ESTADO CIVIL"].value_counts().reset_index()
    df_grouped.columns = ["Estado Civil", "Cantidad"]
    fig = go.Figure(data=[go.Pie(labels=df_grouped["Estado Civil"], values=df_grouped["Cantidad"], hole=0.3)])
    fig.update_traces(pull=[0.1] * len(df_grouped), textinfo="percent", marker=dict(line=dict(color="black", width=1)))
    fig.update_layout(
        title={"text": "Estado Civil", "x": 0.5, "xanchor": "center"},
        scene_camera=dict(eye=dict(x=5, y=2, z=0.3)),
        width=300,
        height=300
    )
    fig.write_image(img_path)

# Gráfico de barras para generación
def grafico_generacion(df, img_path):
    def categorizar_generacion(ano):
        if 1946 <= ano <= 1964:
            return "Baby Boomers"
        elif 1965 <= ano <= 1980:
            return "Generación X"
        elif 1981 <= ano <= 1994:
            return "Millennials"
        
        elif 1995 <= ano <= 2000:
            return "Generación Z"
        else:
            return "Otras Generaciones"
    df["Generación"] = df["3.AÑO DE NACIMIENTO"].apply(categorizar_generacion)
    conteo_generaciones = df["Generación"].value_counts().reset_index()
    conteo_generaciones.columns = ["Generación", "Cantidad"]
    fig = px.bar(
        conteo_generaciones,
        x="Generación",
        y="Cantidad",
        title="Distribución por Generaciones",
        labels={"Cantidad": "Cantidad de Personas", "Generación": "Generación"},
        color="Generación",
        width=500,
        height=300
    )
    fig.write_image(img_path)

# Gráfico de barras para escolaridad
def grafico_escolaridad(df, img_path):
    df_grouped = df["5. NIVEL DE ESTUDIOS"].value_counts()
    total = df_grouped.sum()
    plt.figure(figsize=(4, 3))
    ax = sns.barplot(
        x=df_grouped.index,
        y=df_grouped.values,
        palette="viridis"
    )
    plt.title("Distribución de Nivel de Estudios")
    plt.xlabel("Nivel de Estudios")
    plt.ylabel("Número de Personas")
    plt.xticks(rotation=10, fontsize=8)
    for i, v in enumerate(df_grouped.values):
        porcentaje = v / total * 100
        ax.text(i, v * 0.5, f"{porcentaje:.1f}%", ha='center', va='bottom', fontsize=9, color='black')
    plt.tight_layout()
    plt.savefig(img_path)
    plt.close()

# =======================
# FUNCIONES DE DOMINIOS
# =======================
# # Diccionario de dominios
Dominio = { 
    "Características del liderazgo Tipo A": {
		"columnas": [
			"Mi jefe me da instrucciones claras",
			"Mi jefe ayuda a organizar mejor el trabajo",
			"Mi jefe tiene en cuenta mis puntos de vista y opiniones",
			"Mi jefe me anima para hacer mejor mi trabajo",
			"Mi jefe distribuye las tareas de forma que me facilita el trabajo",
			"Mi jefe me comunica a tiempo la información relacionada con el trabajo",
			"La orientación que me da mi jefe me ayuda a hacer mejor el trabajo",
			"Mi jefe me ayuda a progresar en el trabajo",
			"Mi jefe me ayuda a sentirme bien en el trabajo",
			"Mi jefe ayuda a solucionar los problemas que se presentan en el trabajo",
			"Siento que puedo confiar en mi jefe",
			"Mi jefe me escucha cuando tengo problemas de trabajo",
			"Mi jefe me brinda su apoyo cuando lo necesito"
		],
		"divisor":52 ,
		"limites_riesgo": {
			"muy_alto": (46.3, 100),
			"alto": (30.9, 46.2),
			"medio": (15.5, 30.8),
			"bajo": (3.9, 15.4),
			"sin_riesgo": (0.0, 3.8),
		}
    },
    "Relaciones sociales en el trabajo Tipo A": {
		"columnas": [
			'Me agrada el ambiente de mi grupo de trabajo',
			'En mi grupo de trabajo me tratan de forma respetuosa',
			'Siento que puedo confiar en mis compañeros de trabajo',
			'Me siento a gusto con mis compañeros de trabajo',
			'En mi grupo de trabajo algunas personas me maltratan',
			'Entre compañeros solucionamos los problemas de forma respetuosa',
			'Hay integración en mi grupo de trabajo',
			'Mi grupo de trabajo es muy unido',
			'Las personas en mi trabajo me hacen sentir parte del grupo',
			'Es fácil poner de acuerdo al grupo para hacer el trabajo',
			'Mis compañeros de trabajo me ayudan cuando tengo dificultades',
			'En mi trabajo las personas nos apoyamos unos a otros',
			'Algunos compañeros de trabajo me escuchan cuando tengo problemas',
			'Cuando tenemos que realizar trabajo de grupo los compañeros colaboran'
		],
		"divisor":56 ,
		"limites_riesgo": {
			"muy_alto": (37.6, 100),
			"alto": (25.1, 37.5),
			"medio": (16.2, 25.0),
			"bajo": (5.5, 16.1),
			"sin_riesgo": (0.0, 5.4),
		}
    },
    "Retroalimentación del desempeño Tipo A": {
		"columnas": [
			'Me informan sobre lo que hago bien en mi trabajo',
			'Me informan sobre lo que debo mejorar en mi trabajo',
			'La información que recibo sobre mi rendimiento en el trabajo es clara',
			'La forma como evalúan mi trabajo en la empresa me ayuda a mejorar',
			'Me informan a tiempo sobre lo que debo mejorar en el trabajo'
		],
		"divisor": 20,
		"limites_riesgo": {
			"muy_alto": (55.1, 100),
			"alto": (40.1, 55.0),
			"medio": (25.1, 40.0),
			"bajo": (10.1, 25.0),
			"sin_riesgo": (0.0, 10.0),
		}
    },
    "Relación con los colaboradores Tipo A": {
		"columnas": [
			'Tengo colaboradores que comunican tarde los asuntos de trabajo',
			'Tengo colaboradores que tienen comportamientos irrespetuosos',
			'Tengo colaboradores que dificultan la organización del trabajo',
			'Tengo colaboradores que guardan silencio cuando les piden opiniones',
			'Tengo colaboradores que dificultan el logro de los resultados del trabajo',
			'Tengo colaboradores que expresan de forma irrespetuosa sus desacuerdos',
			'Tengo colaboradores que cooperan poco cuando se necesita',
			'Tengo colaboradores que me preocupan por su desempeño',
			'Tengo colaboradores que ignoran las sugerencias para mejorar su trabajo'
		],
		"divisor": 36,
		"limites_riesgo": {
			"muy_alto": (45.3, 100),
			"alto": (33.4, 47.2),
			"medio": (25.1, 33.3),
			"bajo": (14.0, 25.0),
			"sin_riesgo": (0.0, 13.9),
		}
    },
    "Características del liderazgo Tipo B": {
		"columnas": [
			"Mi jefe me da instrucciones claras",
			"Mi jefe ayuda a organizar mejor el trabajo",
			"Mi jefe tiene en cuenta mis puntos de vista y opiniones",
			"Mi jefe me anima para hacer mejor mi trabajo",
			"Mi jefe distribuye las tareas de forma que me facilita el trabajo",
			"Mi jefe me comunica a tiempo la información relacionada con el trabajo",
			"La orientación que me da mi jefe me ayuda a hacer mejor el trabajo",
			"Mi jefe me ayuda a progresar en el trabajo",
			"Mi jefe me ayuda a sentirme bien en el trabajo",
			"Mi jefe ayuda a solucionar los problemas que se presentan en el trabajo",
			"Siento que puedo confiar en mi jefe",
			"Mi jefe me escucha cuando tengo problemas de trabajo",
			"Mi jefe me brinda su apoyo cuando lo necesito"
		],
		"divisor": 52,
		"limites_riesgo": {
			"muy_alto": (38.6, 100),
			"alto": (25.1, 38.5),
			"medio": (13.6, 25.0),
			"bajo": (3.9, 13.5),
			"sin_riesgo": (0.0, 3.8),
		}
    },
    "Relaciones sociales en el trabajo Tipo B": {
		"columnas": [
			'Me agrada el ambiente de mi grupo de trabajo',
			'En mi grupo de trabajo me tratan de forma respetuosa',
			'Siento que puedo confiar en mis compañeros de trabajo',
			'Me siento a gusto con mis compañeros de trabajo',
			'En mi grupo de trabajo algunas personas me maltratan',
			'Entre compañeros solucionamos los problemas de forma respetuosa',
			'Hay integración en mi grupo de trabajo',
			'Mi grupo de trabajo es muy unido',
			'Las personas en mi trabajo me hacen sentir parte del grupo',
			'Es fácil poner de acuerdo al grupo para hacer el trabajo',
			'Mis compañeros de trabajo me ayudan cuando tengo dificultades',
			'En mi trabajo las personas nos apoyamos unos a otros',
			'Algunos compañeros de trabajo me escuchan cuando tengo problemas',
			'Cuando tenemos que realizar trabajo de grupo los compañeros colaboran'
		],
		"divisor": 56,
		"limites_riesgo": {
			"muy_alto": (37.6, 100),
			"alto": (27.2, 37.5),
			"medio": (14.7, 27.1),
			"bajo": (6.4, 14.6),
			"sin_riesgo": (0.0, 6.3),
		}
    },
    "Retroalimentación del desempeño Tipo B": {
		"columnas": [
			'Me informan sobre lo que hago bien en mi trabajo',
			'Me informan sobre lo que debo mejorar en mi trabajo',
			'La información que recibo sobre mi rendimiento en el trabajo es clara',
			'La forma como evalúan mi trabajo en la empresa me ayuda a mejorar',
			'Me informan a tiempo sobre lo que debo mejorar en el trabajo'
		],
		"divisor": 20,
		"limites_riesgo": {
			"muy_alto": (50.1, 100),
			"alto": (30.1, 50.0),
			"medio": (20.1, 30.0),
			"bajo": (5.1, 20.0),
			"sin_riesgo": (0.0, 5.0),
		}
    },	
    "Capacitación Tipo A": {
		"columnas": [
			'Me informan con claridad cuáles son mis funciones',
			'Me informan cuáles son las decisiones que puedo tomar en mi trabajo',
			'Me explican claramente los resultados que debo lograr en mi trabajo',
			'Me explican claramente el efecto de mi trabajo en la empresa',
			'Me explican claramente los objetivos de mi trabajo',
			'Me informan claramente quien me puede orientar para hacer mi trabajo',
			'Me informan claramente con quien puedo resolver los asuntos de trabajo'
		],
		"divisor": 12,
		"limites_riesgo": {
			"muy_alto": (50.1, 100),
			"alto": (33.4, 50.0),
			"medio": (16.8, 33.3),
			"bajo": (1.0, 16.7),
			"sin_riesgo": (0.0, 0.9),
		}
    },
    "Claridad de rol Tipo A": {
		"columnas": [
			'Me informan con claridad cuáles son mis funciones',
			'Me informan cuáles son las decisiones que puedo tomar en mi trabajo',
			'Me explican claramente los resultados que debo lograr en mi trabajo',
			'Me explican claramente el efecto de mi trabajo en la empresa',
			'Me explican claramente los objetivos de mi trabajo',
			'Me informan claramente quien me puede orientar para hacer mi trabajo',
			'Me informan claramente con quien puedo resolver los asuntos de trabajo'
		],
		"divisor": 28,
		"limites_riesgo": {
			"muy_alto": (39.4, 100),
			"alto": (21.5, 39.3),
			"medio": (10.8, 21.4),
			"bajo": (1.0, 10.7),
			"sin_riesgo": (0.0, 0.9),
		}
    },
    "Control y autonomía sobre el trabajo Tipo A": {
		"columnas": [
			'Puedo decidir cuánto trabajo hago en el día',
			'Puedo decidir la velocidad a la que trabajo',
			'Puedo cambiar el orden de las actividades en mi trabajo'
		],
		"divisor": 12,
		"limites_riesgo": {
			"muy_alto": (58.4, 100),
			"alto": (41.8, 58.3),
			"medio": (25.1, 41.7),
			"bajo": (8.4, 25.0),
			"sin_riesgo": (0.0, 8.3),
		}
    },
    "Oportunidades para el uso y desarrollo de habilidades y conocimientos Tipo A": {
		"columnas": [
			'Mi trabajo me permite desarrollar mis habilidades',
			'Mi trabajo me permite aplicar mis conocimientos',
			'Mi trabajo me permite aprender nuevas cosas',
			'Me asignan el trabajo teniendo en cuenta mis capacidades.'
		],
		"divisor": 16,
		"limites_riesgo": {
			"muy_alto": (31.4, 100),
			"alto": (18.9, 31.3),
			"medio": (6.4, 18.8),
			"bajo": (1.0, 6.3),
			"sin_riesgo": (0.0, 0.9),
		}
    },	
    "Participación y manejo del cambio Tipo A": {
		"columnas": [
			'Los cambios en mi trabajo han sido beneficiosos',
			'Me explican claramente los cambios que ocurren en mi trabajo',
			'Puedo dar sugerencias sobre los cambios que ocurren en mi trabajo',
			'Cuando se presentan cambios en mi trabajo se tienen en cuenta mis ideas y sugerencias'
		],
		"divisor":16 ,
		"limites_riesgo": {
			"muy_alto": (50.1, 100),
			"alto": (37.6, 50.0),
			"medio": (25.1, 37.5),
			"bajo": (12.6, 25.0),
			"sin_riesgo": (0.0, 12.5),
		}
    },
    "Consistencia del rol Tipo A": {
		"columnas": [
			'En el trabajo me dan órdenes contradictorias',
			'En mi trabajo me piden hacer cosas innecesarias',
			'En mi trabajo se presentan situaciones en las que debo pasar por alto normas o procedimientos',
			'En mi trabajo tengo que hacer cosas que se podrían hacer de una forma más práctica',
			'Los cambios que se presentan en mi trabajo dificultan mi labor'
		],
		"divisor":20,
		"limites_riesgo": {
			"muy_alto": (45.1, 100),
			"alto": (35.1, 45.0),
			"medio": (25.1, 35.0),
			"bajo": (15.1, 25.0),
			"sin_riesgo": (0.0, 15.0),
		}
    },
    "Demandas ambientales y de esfuerzo físico Tipo A": {
		"columnas": [
			'El ruido en el lugar donde trabajo es molesto',
			'En el lugar donde trabajo hace mucho frío',
			'En el lugar donde trabajo hace mucho calor',
			'El aire en el lugar donde trabajo es fresco y agradable',
			'La luz del sitio donde trabajo es agradable',
			'El espacio donde trabajo es cómodo',
			'En mi trabajo me preocupa estar expuesto asustancias químicas que afecten mi salud',
			'Mi trabajo me exige hacer mucho esfuerzo físico',
			'Los equipos o herramientas con los que trabajo son cómodos',
			'En mi trabajo me preocupa estar expuesto a microbios, animales o plantas que afecten mi salud',
			'Me preocupa accidentarme en mi trabajo',
			'El lugar donde trabajo es limpio y ordenado'
		],
		"divisor": 48,
		"limites_riesgo": {
			"muy_alto": (39.7, 100),
			"alto": (31.4, 39.6),
			"medio": (23.0, 31.3),
			"bajo": (14.7, 22.9),
			"sin_riesgo": (0.0, 14.6),
		}
    },
    "Demandas cuantitativas Tipo A": {
		"columnas": [
			'Por la cantidad de trabajo que tengo debo quedarme tiempo adicional',
			'Me alcanza el tiempo de trabajo para tener al día mis deberes',
			'Por la cantidad de trabajo que tengo debo trabajar sin parar',
			'En mi trabajo es posible tomar pausas para descansar',
			'Puedo tomar pausas cuando las necesito',
			'Puedo parar un momento mi trabajo para atender algún asunto personal'
		],
		"divisor": 24,
		"limites_riesgo": {
			"muy_alto": (54.3, 100),
			"alto": (45.9, 54.2),
			"medio": (33.4, 45.8),
			"bajo": (25.1, 33.3),
			"sin_riesgo": (0.0, 25.0),
		}
    },	
    "Demandas de carga mental Tipo A": {
		"columnas": [
			'Mi trabajo me exige hacer mucho esfuerzo mental',
			'Mi trabajo me exige estar muy concentrado',
			'Mi trabajo me exige memorizar mucha información',
			'Mi trabajo me exige atender a muchos asuntos al mismo tiempo',
			'Mi trabajo requiere que me fije en pequeños detalles'
		],
		"divisor": 20,
		"limites_riesgo": {
			"muy_alto": (54.3, 100),
			"alto": (45.9, 54.2),
			"medio": (33.4, 45.8),
			"bajo": (25.1, 33.3),
			"sin_riesgo": (0.0, 25.0),
		}
    },
    "Demandas de la jornada de trabajo Tipo A": {
		"columnas": [
			'Trabajo en horario de noche',
			'Mi trabajo me exige laborar en días de descanso, festivos o fines de semana',
			'En mi trabajo puedo tomar fines de semana o días de descanso al mes'
		],
		"divisor": 12,
		"limites_riesgo": {
			"muy_alto": (50.1, 100),
			"alto": (33.4, 50.0),
			"medio": (25.1, 33.3),
			"bajo": (8.4, 25.0),
			"sin_riesgo": (0.0, 8.3),
		}
    },
    "Demandas emocionales Tipo A": {
		"columnas": [
			'Atiendo clientes o usuarios muy enojados',
			'Atiendo clientes o usuarios muy preocupados',
			'Atiendo clientes o usuarios muy tristes',
			'Mi trabajo me exige atender personas muy enfermas',
			'Mi trabajo me exige atender personas muy necesitadas de ayuda',
			'Atiendo clientes o usuarios que me maltratan',
			'Para hacer mi trabajo debo demostrar sentimientos distintos a los míos',
			'Mi trabajo me exige atender situaciones de violencia',
			'Mi trabajo me exige atender situaciones muy tristes o dolorosas'
		],
		"divisor": 36,
		"limites_riesgo": {
			"muy_alto": (47.3, 100),
			"alto": (33.4, 47.2),
			"medio": (25.1, 33.3),
			"bajo": (16.8, 25.0),
			"sin_riesgo": (0.0, 16.7),
		}
    },
    "Influencia del trabajo sobre el entorno extralaboral Tipo A": {
		"columnas": [
			'Cuando estoy en casa sigo pensando en el trabajo',
			'Discuto con mi familia o amigos por causa de mi trabajo',
			'Debo atender asuntos de trabajo cuando estoy en casa',
			'Por mi trabajo el tiempo que paso con mi familia y amigos es muy poco'
		],
		"divisor": 16,
		"limites_riesgo": {
			"muy_alto": (50.1, 100),
			"alto": (43.9, 50.0),
			"medio": (31.4, 43.8),
			"bajo": (18.9, 31.3),
			"sin_riesgo": (0.0, 18.8),
		}
    },	
    "Nivel de responsabilidad del cargo Tipo A": {
		"columnas": [
			'En mi trabajo tengo que tomar decisiones difíciles muy rápido',
			'En mi trabajo respondo por cosas de mucho valor',
			'En mi trabajo respondo por dinero de la empresa',
			'Como parte de mis funciones debo responder por la seguridad de otros',
			'Respondo ante mi jefe por los resultados de toda mi área de trabajo',
			'Mi trabajo me exige cuidar la salud de otras personas'
		],
		"divisor": 24,
		"limites_riesgo": {
			"muy_alto": (79.3, 100),
			"alto": (66.8, 79.2),
			"medio": (54.3, 66.7),
			"bajo": (37.6, 54.2),
			"sin_riesgo": (0.0, 37.5),
		}
    },
    "Recompensas derivadas de la pertenencia a la organización y del trabajo que se realiza Tipo A": {
       "columnas": [
           'Mi trabajo en la empresa es estable',
           'El trabajo que hago me hace sentir bien',
           'Siento orgullo de trabajar en esta empresa',
           'Hablo bien de la empresa con otras personas'
       ],
       "divisor": 16,
       "limites_riesgo": {
           "muy_alto": (18.9, 100),
           "alto": (12.6, 18.8),
           "medio": (6.4, 12.5),
           "bajo": (1.0, 6.3),
           "sin_riesgo": (0.0, 0.9)
       }
   },
   "Reconocimiento y compensación Tipo A": {
       "columnas": [
           'En la empresa me pagan a tiempo mi salario',
           'El pago que recibo es el que me ofreció la empresa',
           'El pago que recibo es el que merezco por el trabajo que realizo',
           'En mi trabajo tengo posibilidades de progresar',
           'Las personas que hacen bien el trabajo pueden progresar en la empresa',
           'La empresa se preocupa por el bienestar de los trabajadores'
       ],
       "divisor": 24,
       "limites_riesgo": {
           "muy_alto": (37.6, 100),
           "alto": (25.1, 37.5),
           "medio": (12.6, 25.0),
           "bajo": (1.0, 12.5),
           "sin_riesgo": (0.0, 0.9)
       }
   },
   "Capacitación Tipo B": {
		"columnas": [
			"La empresa me permite asistir a capacitaciones relacionadas con mi trabajo",
			"Recibo capacitación útil para hacer mi trabajo",
			"Recibo capacitación que me ayuda a hacer mejor mi trabajo"
		],
		"divisor": 12,
		"limites_riesgo": {
			"muy_alto": (50.1, 100),
			"alto": (25.1, 50.0),
			"medio": (16.8, 25.0),
			"bajo": (1.0, 16.7),
			"sin_riesgo": (0.0, 0.9),
		}
    },
    "Claridad de rol Tipo B": {
		"columnas": [
			"Me informan con claridad cuáles son mis funciones",
			"Me informan cuáles son las decisiones que puedo tomar en mi trabajo",
			"Me explican claramente los resultados que debo lograr en mi trabajo",
			"Me explican claramente los objetivos de mi trabajo",
			"Me informan claramente con quien puedo resolver los asuntos de trabajo"
		],
		"divisor": 20,
		"limites_riesgo": {
			"muy_alto": (30.1, 100),
			"alto": (15.1, 30.0),
			"medio": (5.1, 15.0),
			"bajo": (1.0, 5.0),
			"sin_riesgo": (0.0, 0.9),
		}
    },
    "Control y autonomía sobre el trabajo Tipo B": {
		"columnas": [
			"Puedo decidir cuánto trabajo hago en el día",
			"Puedo decidir la velocidad a la que trabajo",
			"Puedo cambiar el orden de las actividades en mi trabajo"
		],
		"divisor": 12,
		"limites_riesgo": {
			"muy_alto": (75.1, 100),
			"alto": (66.8, 75.0),
			"medio": (50.1, 66.7),
			"bajo": (33.4, 50.0),
			"sin_riesgo": (0.0, 33.3),
		}
    },	
    "Oportunidades para el uso y desarrollo de habilidades y conocimientos Tipo B": {
		"columnas": [
			"En mi trabajo puedo hacer cosas nuevas",
			"Mi trabajo me permite desarrollar mis habilidades",
			"Mi trabajo me permite aplicar mis conocimientos",
			"Mi trabajo me permite aprender nuevas cosas"
		],
		"divisor": 16,
		"limites_riesgo": {
			"muy_alto": (56.4, 100),
			"alto": (37.6, 56.3),
			"medio": (25.1, 37.5),
			"bajo": (12.6, 25.0),
			"sin_riesgo": (0.0, 12.5),
		}
    },
    "Participación y manejo del cambio Tipo B": {
		"columnas": [
			"Me explican claramente los cambios que ocurren en mi trabajo",
			"Puedo dar sugerencias sobre los cambios que ocurren en mi trabajo",
			"Cuando se presentan cambios en mi trabajo se tienen en cuenta mis ideas y sugerencias"
		],
		"divisor": 12,
		"limites_riesgo": {
			"muy_alto": (75.1, 100),
			"alto": (66.8, 75.0),
			"medio": (50.1, 66.7),
			"bajo": (33.4, 50.0),
			"sin_riesgo": (0.0, 33.3),
		}
    },
    "Demandas ambientales y de esfuerzo físico Tipo B": {
		"columnas": [
			'El ruido en el lugar donde trabajo es molesto',
			'En el lugar donde trabajo hace mucho frío',
			'En el lugar donde trabajo hace mucho calor',
			'El aire en el lugar donde trabajo es fresco y agradable',
			'La luz del sitio donde trabajo es agradable',
			'El espacio donde trabajo es cómodo',
			'En mi trabajo me preocupa estar expuesto asustancias químicas que afecten mi salud',
			'Mi trabajo me exige hacer mucho esfuerzo físico',
			'Los equipos o herramientas con los que trabajo son cómodos',
			'En mi trabajo me preocupa estar expuesto a microbios, animales o plantas que afecten mi salud',
			'Me preocupa accidentarme en mi trabajo12',
			'El lugar donde trabajo es limpio y ordenado13'
		],
		"divisor": 48,
		"limites_riesgo": {
			"muy_alto": (48.0, 100),
			"alto": (39.7, 47.9),
			"medio": (31.4, 39.6),
			"bajo": (23.0, 31.3),
			"sin_riesgo": (0.0, 22.9),
		}
    },
    "Demandas cuantitativas Tipo B": {
		"columnas": [
			'Por la cantidad de trabajo que tengo debo quedarme tiempo adicional',
			'Me alcanza el tiempo de trabajo para tener al día mis deberes',
			'Por la cantidad de trabajo que tengo debo trabajar sin parar'
		],
		"divisor": 12,
		"limites_riesgo": {
			"muy_alto": (50.1, 100),
			"alto": (41.8, 50.0),
			"medio": (33.4, 41.7),
			"bajo": (16.8, 33.3),
			"sin_riesgo": (0.0, 16.7),
		}
    },
    "Demandas de carga mental Tipo B": {
		"columnas": [
			'Mi trabajo me exige hacer mucho esfuerzo mental',
			'Mi trabajo me exige estar muy concentrado18',
			'Mi trabajo me exige memorizar mucha información',
			'En mi trabajo tengo que hacer cálculos matemáticos',
			'Mi trabajo requiere que me fije en pequeños detalles'
		],
		"divisor": 20,
		"limites_riesgo": {
			"muy_alto": (85.1, 100),
			"alto": (75.1, 85.0),
			"medio": (65.1, 75.0),
			"bajo": (50.1, 65.0),
			"sin_riesgo": (0.0, 50.0),
		}
    },
    "Demandas de la jornada de trabajo Tipo B": {
		"columnas": [
			'Trabajo en horario de noche',
			'En mi trabajo es posible tomar pausas para descansar',
			'Mi trabajo me exige laborar en días de descanso, festivos o fines de semana',
			'En mi trabajo puedo tomar fines de semana o días de descanso al mes',
			'Puedo tomar pausas cuando las necesito',
			'Puedo parar un momento mi trabajo para atender algún asunto personal'
		],
		"divisor": 24,
		"limites_riesgo": {
			"muy_alto": (47.3, 100),
			"alto": (39.0, 47.2),
			"medio": (27.9, 38.9),
			"bajo": (19.5, 27.8),
			"sin_riesgo": (0.0, 19.4),
		}
    },
    "Demandas emocionales Tipo B": {
		"columnas": [
			'Atiendo clientes o usuarios muy enojados',
			'Atiendo clientes o usuarios muy preocupados',
			'Atiendo clientes o usuarios muy tristes',
			'Mi trabajo me exige atender personas muy enfermas',
			'Mi trabajo me exige atender personas muy necesitadas de ayuda',
			'Atiendo clientes o usuarios que me maltratan',
			'Mi trabajo me exige atender situaciones de violencia',
			'Mi trabajo me exige atender situaciones muy tristes o dolorosas',
			'Puedo expresar tristeza o enojo frente a las personas que atiendo'
		],
		"divisor": 36,
		"limites_riesgo": {
			"muy_alto": (47.3, 100),
			"alto": (39.0, 47.2),
			"medio": (27.9, 38.9),
			"bajo": (19.5, 27.8),
			"sin_riesgo": (0.0, 19.4),
		}
    },
    "Influencia del trabajo sobre el entorno extralaboral Tipo B": {
		"columnas": [
			'Cuando estoy en casa sigo pensando en el trabajo',
			'Discuto con mi familia o amigos por causa de mi trabajo',
			'Debo atender asuntos de trabajo cuando estoy en casa',
			'Por mi trabajo el tiempo que paso con mi familia y amigos es muy poco'
		],
		"divisor": 16,
		"limites_riesgo": {
			"muy_alto": (50.1, 100),
			"alto": (31.4, 50.0),
			"medio": (25.1, 31.3),
			"bajo": (12.6, 25.0),
			"sin_riesgo": (0.0, 12.5),
		}
    },
    "Recompensas derivadas de la pertenencia a la organización y del trabajo que se realiza Tipo B": {
		"columnas": [
			'Mi trabajo en la empresa es estable',
			'El trabajo que hago me hace sentir bien',
			'Siento orgullo de trabajar en esta empresa',
			'Hablo bien de la empresa con otras personas'
		],
		"divisor": 16,
		"limites_riesgo": {
			"muy_alto": (18.9, 100),
			"alto": (12.6, 18.8),
			"medio": (6.4, 12.5),
			"bajo": (1.0, 6.3),
			"sin_riesgo": (0.0, 0.9),
		}
    },
    "Reconocimiento y compensación Tipo B": {
		"columnas": [
			'En la empresa me pagan a tiempo mi salario',
			'El pago que recibo es el que me ofreció la empresa',
			'El pago que recibo es el que merezco por el trabajo que realizo',
			'En mi trabajo tengo posibilidades de progresar',
			'Las personas que hacen bien el trabajo pueden progresar en la empresa',
			'La empresa se preocupa por el bienestar de los trabajadores'
		],
		"divisor": 24,
		"limites_riesgo": {
			"muy_alto": (37.6, 100),
			"alto": (25.1, 37.5),
			"medio": (12.6, 25.0),
			"bajo": (1.0, 12.5),
			"sin_riesgo": (0.0, 0.9),
		}
    },
        "Tiempo fuera del trabajo": {
        "columnas": [
            'Me queda tiempo para actividades de recreación',
            'Fuera del trabajo tengo tiempo suficiente para descansar',
            'Tengo tiempo para atender mis asuntos personales y del hogar',
            'Tengo tiempo para compartir con mi familia o amigos'
        ],
        "divisor": 16,
        "limites_riesgo": {
            "muy_alto": (50.1, 100),
            "alto": (37.6, 50.0),
            "medio": (25.1, 37.5),
            "bajo": (6.4, 25.0),
            "sin_riesgo": (0.0, 6.3)
        }
    },
        "Relaciones familiares": {
        "columnas": [
            'Cuento con el apoyo de mi familia cuando tengo problemas',
            'La relación con mi familia cercana es cordial',
            'Los problemas con mis familiares los resolvemos de manera amistosa',
        ],
        "divisor": 12,
        "limites_riesgo": {
            "muy_alto": (50.1, 100),
            "alto": (33.4, 50.0),
            "medio": (25.1, 33.3),
            "bajo": (8.4, 25.0),
            "sin_riesgo": (0.0, 8.3)
        }
    },
       "Comunicación y relaciones interpersonales": {
        "columnas": [
            'Tengo buena comunicación con las personas cercanas',
            'Las relaciones con mis amigos son buenas',
            'Converso con personas cercanas sobre diferentes temas',
            'Mis amigos están dispuestos a escucharme cuando tengo problemas',
            'Puedo hablar con personas cercanas sobre las cosas que me pasan'
        ],
        "divisor": 20,
        "limites_riesgo": {
            "muy_alto": (30.1, 100),
            "alto": (20.1, 30.0),
            "medio": (10.1, 20.0),
            "bajo": (1.0, 10.0),
            "sin_riesgo": (0.0, 0.9)
        }
    },
     "Situación económica": {
        "columnas": [
            'El dinero que ganamos en el hogar alcanza para cubrir los gastos básicos',
            'Las relaciones con mis amigos son buenas',
            'Tengo otros compromisos económicos que afectan mucho el presupuesto familiar',
        ],
        "divisor": 12,
        "limites_riesgo": {
            "muy_alto": (50.1, 100),
            "alto": (33.4, 50.0),
            "medio": (25.1, 33.3),
            "bajo": (8.4, 25.0),
            "sin_riesgo": (0.0, 8.3)
        }
    },
  
     "Características de la vivienda y de su entorno": {
        "columnas": [
            'La zona donde vivo es segura',
            'En la zona donde vivo se presentan hurtos y mucha delincuencia',
            'Desde donde vivo me es fácil llegar al centro médico donde me atienden',
            'Cerca a mi vivienda las vías están en buenas condiciones',
            'Cerca a mi vivienda encuentro fácilmente transporte',
            'Las condiciones de mi vivienda son buenas',
            'En mi vivienda hay servicios de agua y luz',
            'Las condiciones de mi vivienda me permiten descansar cuando lo requiero',
            'Las condiciones de mi vivienda me permiten sentirme cómodo'
        ],
        "divisor": 36,
        "limites_riesgo": {
            "muy_alto": (22.3, 100),
            "alto": (14.0, 22.2),
            "medio": (11.2, 13.9),
            "bajo": (5.7, 11.1),
            "sin_riesgo": (0.0, 5.6)
        }
    },
    "Influencia del entorno extralaboral sobre el trabajo":{
        "columnas": [   
            'Mis problemas personales o familiares afectan mi trabajo',
            'Mis problemas personales o familiares me quitan la energía que necesito para trabajar',
            'Mis problemas personales o familiares afectan mis relaciones en el trabajo',
            ],
        "divisor": 12,
        "limites_riesgo":{
            "muy_alto": (41.8, 100),
            "alto": (25.1, 41.7),
            "medio": (16.8, 25.0),
            "bajo": (8.4, 16.7),
            "sin_riesgo": (0.0, 8.3),
        }

    },
    "Desplazamiento vivienda - trabajo - vivienda":{
        "columnas": [
        ],
        "divisor": 16,
        "limites_riesgo": {     
            "muy_alto": (43.9, 100),
            "alto": (25.1, 43.8),
            "medio": (12.6, 25.0),
            "bajo": (1.0, 12.5),
            "sin_riesgo": (0.0, 0.9),
        }

    }


}

# Función para calcular una transformación
def calcular_dominio(df, nombre_dominio):
    if nombre_dominio not in Dominio:
        raise ValueError(f"El Dominio '{nombre_dominio}' no está definido.")
    config = Dominio[nombre_dominio]
    columnas = config["columnas"]
    divisor = config["divisor"]

    if nombre_dominio not in df.columns:
        df[nombre_dominio] = (
            df[columnas].apply(
                lambda row: row.sum() if not row.isnull().any() else None, axis=1
            ) / divisor * 100
        )
    return df

# Función para clasificar riesgos
def clasificar_riesgos(valor, limites):
    if pd.isnull(valor):
        return None
    equivalencias = {
        "muy_alto": "Riesgo muy alto",
        "alto": "Riesgo alto",
        "medio": "Riesgo medio",
        "bajo": "Riesgo bajo",
        "sin_riesgo": "Sin riesgo"
    }
    for nivel, (min_val, max_val) in limites.items():
        if min_val <= valor <= max_val:
            return equivalencias[nivel]
    return "No clasificado"

# funciones de graficos tipo A

def grafico_liderazgo_y_relaciones_sociales_tipo_A(df, img_path):
    dominios = [
        "Características del liderazgo Tipo A",
        "Relaciones sociales en el trabajo Tipo A",
        "Retroalimentación del desempeño Tipo A",
        "Relación con los colaboradores Tipo A"
    ]
    titulo = "Liderazgo y Relaciones Sociales Tipo A"
    crear_grafico(df, dominios, titulo, img_path)

def grafico_control_y_autonomia_tipo_A(df, img_path):
    dominios = [
        "Capacitación Tipo A",
        "Claridad de rol Tipo A",
        "Control y autonomía sobre el trabajo Tipo A",
        "Oportunidades para el uso y desarrollo de habilidades y conocimientos Tipo A",
        "Participación y manejo del cambio Tipo A"
    ]
    titulo = "Control y Autonomía Tipo A"
    crear_grafico(df, dominios, titulo, img_path)

def grafico_demandas_del_trabajo_tipo_A(df, img_path):
    dominios = [
        "Consistencia del rol Tipo A",
        "Demandas ambientales y de esfuerzo físico Tipo A",
        "Demandas cuantitativas Tipo A",
        "Demandas de carga mental Tipo A",
        "Demandas de la jornada de trabajo Tipo A",
        "Demandas emocionales Tipo A",
        "Influencia del trabajo sobre el entorno extralaboral Tipo A",
        "Nivel de responsabilidad del cargo Tipo A"
    ]
    titulo = "Demandas del trabajo Tipo A"
    crear_grafico(df, dominios, titulo, img_path)

def recompensas_tipo_A(df, img_path):
    dominios = [
        "Recompensas derivadas de la pertenencia a la organización y del trabajo que se realiza Tipo A",
        "Reconocimiento y compensación Tipo A"
    ]
    titulo = "Recompensas Tipo A"
    crear_grafico(df, dominios, titulo, img_path)

# funciones de graficos tipo B

def grafico_liderazgo_y_relaciones_sociales_tipo_B(df, img_path):
    dominios = [
        "Características del liderazgo Tipo B",
        "Relaciones sociales en el trabajo Tipo B",
        "Retroalimentación del desempeño Tipo B"
    ]
    titulo = "Lidrazgo y Relaciones Sociales Tipo B"
    crear_grafico(df, dominios, titulo, img_path)

def control_y_autonomia_tipo_B(df, img_path):
    dominios = [
        "Capacitación Tipo B",
        "Claridad de rol Tipo B",
        "Control y autonomía sobre el trabajo Tipo B",
        "Oportunidades para el uso y desarrollo de habilidades y conocimientos Tipo B",
        "Participación y manejo del cambio Tipo B"
    ]
    titulo = "Control y Autonomía Tipo B"
    crear_grafico(df, dominios, titulo, img_path)

def demandas_de_trabajo_tipo_B(df, img_path):
    dominios = [
        "Demandas ambientales y de esfuerzo físico Tipo B",
        "Demandas cuantitativas Tipo B",
        "Demandas de carga mental Tipo B",
        "Demandas de la jornada de trabajo Tipo B",
        "Demandas emocionales Tipo B",
        "Influencia del trabajo sobre el entorno extralaboral Tipo B"
    ]
    titulo = "Demandas del trabajo Tipo B"
    crear_grafico(df, dominios, titulo, img_path)

def grafico_recompensas_tipo_B(df, img_path):
    dominios = [
        "Recompensas derivadas de la pertenencia a la organización y del trabajo que se realiza Tipo B",
        "Reconocimiento y compensación Tipo B"
    ]
    titulo = "Recompensas Tipo B"
    crear_grafico(df, dominios, titulo, img_path)

# funcion de graficos extralaborales

def factores_riesgos_extralaborales(df, img_path):
    dominios = [
        "Tiempo fuera del trabajo",
        "Relaciones familiares",
        "Comunicación y relaciones interpersonales",
        "Situación económica",
        "Características de la vivienda y de su entorno",
        "Influencia del entorno extralaboral sobre el trabajo",
        "Desplazamiento vivienda - trabajo - vivienda"
    ]
    titulo = "Factores de Riesgo Extralaborales"
    crear_grafico(df, dominios, titulo, img_path)


# =======================
# FUNCIONES DE GRAFICOS COMBINADOS
# =======================


def crear_grafico(df, dominios, titulo, img_path):
    import textwrap

    niveles = ["Riesgo muy alto", "Riesgo alto", "Riesgo medio", "Riesgo bajo", "Sin riesgo"]
    conteo = []
    max_dim_len = 12  # Máximo de caracteres por línea en la etiqueta de dimensión

    def wrap_label(label):
        return "\n".join(textwrap.wrap(label, max_dim_len))

    dominios_wrapped = [wrap_label(d) for d in dominios]

    for nombre_dominio, nombre_dominio_wrapped in zip(dominios, dominios_wrapped):
        if nombre_dominio not in Dominio:
            raise ValueError(f"El Dominio '{nombre_dominio}' no está definido.")
        config = Dominio[nombre_dominio]
        df = calcular_dominio(df, nombre_dominio)

        clasificaciones = df[nombre_dominio].dropna().apply(lambda x: clasificar_riesgos(x, config["limites_riesgo"]))
        total = len(clasificaciones)
        for nivel in niveles:
            cantidad = (clasificaciones == nivel).sum()
            porcentaje = (cantidad / total * 100) if total > 0 else 0
            conteo.append({"Dimensión": nombre_dominio_wrapped, "Riesgo": nivel, "Porcentaje": porcentaje})

    df_plot = pd.DataFrame(conteo)

    ancho = max(8, len(dominios) * 1.1)
    alto = 5
    plt.figure(figsize=(ancho, alto))
    colores = ["#d32f2f", "#fbc02d", "#9b9b9b", "#1976d2", "#388e3c"]
    ax = sns.barplot(
        x="Dimensión", y="Porcentaje", hue="Riesgo",
        data=df_plot, palette=colores, order=dominios_wrapped, hue_order=niveles, dodge=True
    )
    for p in ax.patches:
        height = p.get_height()
        ax.annotate(f'{height:.1f}%',
                    (p.get_x() + p.get_width() / 2., height + 0.5),
                    ha='center', va='bottom', fontsize=8, color='black', rotation=0)
    plt.title(titulo, fontsize=12)
    plt.xlabel("Dimensión", fontsize=11)
    plt.ylabel("Porcentaje (%)", fontsize=11)
    plt.xticks(rotation=0, fontsize=9, ha='center', wrap=True)
    plt.yticks(fontsize=9)
    # Coloca la leyenda fuera del gráfico, horizontal y con espacio a la derecha
    plt.subplots_adjust(right=0.75)
    legend = plt.legend(
        title="Nivel de Riesgo",
        fontsize=8,
        title_fontsize=8,
        loc='center left',
        bbox_to_anchor=(1.01, 0.5),
        borderaxespad=0,
        frameon=True,
        ncol=1  # Una sola columna,  cambiar a ncol=2  horizontal
    )
    plt.tight_layout()
    plt.savefig(img_path)
    plt.close()

# =======================
# FUNCIONES DE INFORME
# =======================

def generar_graficos_y_pdf(file_path, filename, reports_dir):
    df = pd.read_excel(file_path)
    graficos = [
        ("Sexo", grafico_tipo_sexo),
        ("Estado civil", grafico_estadocivil),
        ("Personas a Cargo", grafico_personas_a_cargo),
        ("Tipo de Vivienda", grafico_tipo_vivienda),
        ("Estrato", grafico_estrato),
        ("Tipo de Contrato", grafico_tipo_contrato),
        ("Escolaridad", grafico_escolaridad),
        ("Tipo de generación", grafico_generacion),
        ("Liderazgo y relaciones sociales tipo A", grafico_liderazgo_y_relaciones_sociales_tipo_A),
        ("Control y autonomía tipo A", grafico_control_y_autonomia_tipo_A),
        ("Demandas del trabajo tipo A", grafico_demandas_del_trabajo_tipo_A),
        ("Recompensas tipo A", recompensas_tipo_A),
        ("Liderazgo y relaciones sociales tipo B", grafico_liderazgo_y_relaciones_sociales_tipo_B),
        ("Control y autonomía tipo B", control_y_autonomia_tipo_B),
        ("Demandas de trabajo tipo B", demandas_de_trabajo_tipo_B),
        ("Recompensas tipo B", grafico_recompensas_tipo_B),
        ("Factores de Riesgo Extralaborales", factores_riesgos_extralaborales)
    ]

    img_paths = []
    for i, (titulo, funcion) in enumerate(graficos):
        img_path = os.path.join(reports_dir, f"{filename}_grafico_{i}.png")
        funcion(df, img_path)
        img_paths.append((titulo, img_path))

    pdf_path = os.path.join(reports_dir, f"{filename}_informe.pdf")
    pdf = FPDF()
    pdf.alias_nb_pages()
    pdf.set_font("Arial", size=12)
    img_width = 160  # Ancho deseado para la imagen
    margen = 10

    for idx, (titulo, img) in enumerate(img_paths):
        pdf.add_page()
        if idx == 0:
            # Portada con logo y título
            logo_path = os.path.join(os.path.dirname(__file__), "..", "static", "logo.jpg")
            pdf.image(logo_path, x=10, y=10, w=40)
            pdf.set_font("Arial", "B", 18)
            pdf.set_xy(0, 60)
            pdf.cell(0, 10, "Informe de Riesgo Psicosocial", ln=True, align='C')
        # Título del gráfico
        pdf.set_font("Arial", "B", 14)
        pdf.set_xy(10, 100)
        pdf.cell(0, 10, titulo, align='C')
        # Imagen centrada
        x_centro = (pdf.w - img_width) / 2
        pdf.image(img, x=x_centro, y=110, w=img_width)

    pdf.output(pdf_path)
    return f"{filename}_informe.pdf"

def generar_informe_en_word(file_path, graficos, filename, reports_dir):
    df = pd.read_excel(file_path)
    graficos = [
        ("Sexo", grafico_tipo_sexo),
        ("Estado Civil", grafico_estadocivil),
        ("Numero de personas a cargo", grafico_personas_a_cargo),
        ("Tipo de vivienda", grafico_tipo_vivienda),
        ("Estrato", grafico_estrato),
        ("Tipo de Contrato", grafico_tipo_contrato),
        ("Escolaridad", grafico_escolaridad),
        ("Tipo de generación ", grafico_generacion),
        ("Liderazgo y relaciones sociales tipo A", grafico_liderazgo_y_relaciones_sociales_tipo_A),
        ("Control y autonomía tipo A", grafico_control_y_autonomia_tipo_A),
        ("Demandas del trabajo tipo A", grafico_demandas_del_trabajo_tipo_A),
        ("Recompensas tipo A", recompensas_tipo_A),
        ("Liderazgo y relaciones sociales tipo B", grafico_liderazgo_y_relaciones_sociales_tipo_B),
        ("Control y autonomía tipo B", control_y_autonomia_tipo_B),
        ("Demandas de trabajo tipo B", demandas_de_trabajo_tipo_B),
        ("recompensas tipo B", grafico_recompensas_tipo_B),
        ("Factores de Riesgo Extralaborales", factores_riesgos_extralaborales)
    ]
    
    img_paths = []
    for i, (titulo, funcion) in enumerate(graficos):
        img_path = os.path.join(reports_dir, f"{filename}_grafico_{i}.png")
        funcion(df, img_path)
        img_paths.append((titulo, img_path))

    doc_path = os.path.join(reports_dir, f"{filename}_informe.docx")
    doc = Document()

    # Portada
    logo_path = os.path.join(os.path.dirname(__file__), "..", "static", "logo.jpg")
    doc.add_picture(logo_path, width=Inches(1.5))

    doc.add_paragraph("Informe de Riesgo Psicosocial", style='Title')
    doc.add_paragraph("\n")  # Espacio

    for idx, (titulo, img) in enumerate(img_paths):
        if idx != 0:
            doc.add_page_break()
        doc.add_heading(titulo, level=2)
        doc.add_picture(img, width=Inches(5.0))

    doc.save(doc_path)
    return f"{filename}_informe.docx"


